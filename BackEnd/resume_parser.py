


import io
from typing import Optional
import docx
from pdfminer.high_level import extract_text as pdfminer_extract_text
import mammoth
import easyocr
import numpy as np
from PIL import Image

def extract_text_from_pdf(file) -> str:
    try:
        if hasattr(file, 'read'):
            file.seek(0)
            text = pdfminer_extract_text(file)
        else:
            text = pdfminer_extract_text(file)
        return text or ""
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""

def extract_text_from_docx(file) -> str:
    try:
        if hasattr(file, 'read'):
            file.seek(0)
            doc = docx.Document(file)
        else:
            doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""

def extract_text_from_doc(file) -> str:
    try:
        if hasattr(file, 'read'):
            file.seek(0)
            result = mammoth.convert_to_text(file)
        else:
            with open(file, "rb") as f:
                result = mammoth.convert_to_text(f)
        return result.value or ""
    except Exception as e:
        print(f"DOC extraction error: {e}")
        return ""


def extract_text_from_image(file) -> str:
    try:
        if hasattr(file, 'read'):
            file.seek(0)
            image = Image.open(file).convert('RGB')
        else:
            image = Image.open(file).convert('RGB')
        # Convert PIL image to numpy array for EasyOCR
        img_np = np.array(image)
        reader = easyocr.Reader(['en'], gpu=False)
        result = reader.readtext(img_np, detail=0)
        text = '\n'.join(result)
        return text or ""
    except Exception as e:
        print(f"Image OCR extraction error: {e}")
        return ""

def extract_text(file, filename: str) -> Optional[str]:
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file)
    elif name.endswith(".docx"):
        return extract_text_from_docx(file)
    elif name.endswith(".doc"):
        return extract_text_from_doc(file)
    elif name.endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff')):
        return extract_text_from_image(file)
    else:
        # Try to read as text file (txt, csv, etc.)
        try:
            if hasattr(file, 'read'):
                file.seek(0)
                content = file.read()
                if isinstance(content, bytes):
                    content = content.decode(errors='ignore')
                return content
            else:
                with open(file, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
        except Exception as e:
            print(f"Generic text extraction error: {e}")
            return None
